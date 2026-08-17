# -*- coding: utf-8 -*-
"""Generate the FWC HRMS interview-prep study guide as a .docx file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
ACCENT2 = RGBColor(0x2E, 0x74, 0xB5)  # lighter blue
GREY = RGBColor(0x59, 0x59, 0x59)
CODEBG = "F2F2F2"

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, sz, col in [('Heading 1', 16, ACCENT), ('Heading 2', 13.5, ACCENT2), ('Heading 3', 12, ACCENT2)]:
    st = doc.styles[lvl]
    st.font.name = 'Calibri'
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def para(text="", style=None, bold=False, italic=False, size=None, color=None, after=6, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if size:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(text, sub=False, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet 2' if sub else 'List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.15)
    for line in text.split('\n'):
        r = p.add_run(line + '\n')
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), CODEBG)
    pPr.append(sh)
    return p


def qa(q, a):
    p = doc.add_paragraph()
    r = p.add_run("Q: " + q)
    r.bold = True
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(2)
    pa = doc.add_paragraph()
    pa.add_run("A: ").bold = True
    pa.add_run(a)
    pa.paragraph_format.space_after = Pt(10)
    pa.paragraph_format.left_indent = Inches(0.1)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], "2E74B5")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================================================ COVER
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("FWC HRMS")
r.bold = True; r.font.size = Pt(34); r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Interview Preparation & Learning Roadmap")
r.font.size = Pt(17); r.font.color.rgb = ACCENT2; r.italic = True
para("Everything you need to learn to confidently build and defend an AI-powered MERN project",
     align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, size=11, after=2)
para("Project: AI-Powered Human Resource Management System (MERN + Google Gemini + Brain.js)",
     align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, size=10, italic=True)

doc.add_paragraph()
# what the project is, quick
para("Your Project in One Sentence", style='Heading 2')
para("A full-stack HRMS where four roles (Admin, Senior Manager, HR Recruiter, Employee) manage "
     "employees, attendance, payroll, performance and leaves — augmented with AI: blind resume "
     "screening, adaptive AI interviews with proctoring, an employee flight-risk neural network, "
     "and an HR-policy RAG chatbot — built on React 19, Node/Express, MongoDB, and Google Gemini.")

para("How to use this guide", style='Heading 3')
bullet("Read top to bottom once to map the whole stack, then drill into weak areas.")
bullet("Every section ends with the exact interview questions you are likely to be asked, with model answers.")
bullet("Anything you cannot explain in your own words is a gap — go learn it before the interview.")

doc.add_page_break()

# ============================================================ 0. ROADMAP
para("0. The Learning Roadmap (Big Picture)", style='Heading 1')
para("To build a project like this from scratch, you need competence across six layers. Master them "
     "in this order — each builds on the previous.")
table(
    ["#", "Layer", "What to learn", "Why it matters here"],
    [
     ["1", "Web fundamentals", "HTTP, REST, JSON, client–server model, status codes", "Everything is a REST API call"],
     ["2", "JavaScript (modern)", "ES6+, async/await, promises, modules, the event loop", "Both frontend and backend are JS"],
     ["3", "Frontend", "React 19, hooks, routing, state, data-fetching, CSS", "The entire SPA"],
     ["4", "Backend", "Node.js, Express, middleware, MVC layering", "The REST API server"],
     ["5", "Database", "MongoDB, Mongoose schemas, modeling, queries", "All persistence"],
     ["6", "AI / ML integration", "LLM APIs, prompts, embeddings, RAG, a neural net", "The differentiator"],
    ],
    widths=[0.4, 1.5, 2.6, 2.3])
para("Cross-cutting skills you also need: authentication & security (JWT, bcrypt, RBAC), "
     "Git/GitHub, environment configuration, and deployment.", italic=True, color=GREY)

doc.add_page_break()

# ============================================================ 1. FUNDAMENTALS
para("1. Web & JavaScript Fundamentals", style='Heading 1')

para("1.1 How the web works", style='Heading 2')
bullet("Client–server model: the browser (client) sends an HTTP request; the server responds.", bold_lead="")
bullet("HTTP methods: GET (read), POST (create), PUT/PATCH (update), DELETE (remove).")
bullet("Status codes: 2xx success, 3xx redirect, 4xx client error (400 bad request, 401 unauthorized, 403 forbidden, 404 not found, 409 conflict), 5xx server error.")
bullet("REST: resources addressed by URLs (/api/employees/:id), actions expressed by HTTP verbs, data exchanged as JSON, stateless requests.")
bullet("CORS: browsers block cross-origin requests unless the server sends Access-Control-Allow-* headers. Your server.js configures CORS explicitly.")

para("1.2 JavaScript you must know cold", style='Heading 2')
bullet("Variables & scope: let/const vs var, block scope, hoisting, closures.")
bullet("Functions: arrow functions, default params, rest/spread (...), destructuring.")
bullet("Arrays: map, filter, reduce, find, forEach, some/every — used everywhere in the frontend.")
bullet("Objects: spread/rest, optional chaining (?.), nullish coalescing (??).")
bullet("Asynchronous JS: callbacks → Promises → async/await; Promise.all for parallelism; try/catch for errors.")
bullet("The event loop: call stack, task queue, microtasks — why JS is single-threaded yet non-blocking.")
bullet("Modules: ESM (import/export) on the frontend vs CommonJS (require/module.exports) on the backend.")

para("Async example (the pattern used across your controllers):", style='Heading 3')
code("// async/await with error handling\n"
     "const getEmployee = async (req, res) => {\n"
     "  try {\n"
     "    const emp = await Employee.findById(req.params.id);\n"
     "    if (!emp) return res.status(404).json({ success: false });\n"
     "    res.json({ success: true, data: emp });\n"
     "  } catch (err) {\n"
     "    res.status(500).json({ success: false, message: err.message });\n"
     "  }\n"
     "};")

qa("Why is JavaScript called single-threaded but still able to handle many requests?",
   "It runs one call stack, but I/O operations (DB, network, file) are handed off to the "
   "environment (libuv in Node) and their callbacks are queued. The event loop picks them up when "
   "the stack is free, so the thread is never blocked waiting on I/O.")
qa("Difference between == and ===?",
   "== does type coercion before comparing; === compares value and type with no coercion. Always "
   "prefer === to avoid surprising bugs.")

doc.add_page_break()

# ============================================================ 2. FRONTEND
para("2. Frontend — React 19 + Vite", style='Heading 1')
para("Your frontend is a Single Page Application (SPA): one HTML file, JavaScript swaps the views "
     "client-side via the router, and data comes from REST calls.")

para("2.1 React core concepts", style='Heading 2')
bullet("Components: reusable functions returning JSX (HTML-in-JS). Your pages and components folders are all components.")
bullet("JSX: syntax that compiles to React.createElement; expressions in {curly braces}.")
bullet("Props: read-only inputs passed from parent to child.")
bullet("State: data that changes over time, managed with useState; changing it re-renders the component.")
bullet("Rendering & the virtual DOM: React diffs a virtual tree and patches only what changed — fast updates.")
bullet("Lifecycle via hooks: useEffect runs side-effects (fetch on mount, subscriptions, cleanup).")
bullet("Keys: stable keys in lists let React track items efficiently.")

para("2.2 Hooks (the heart of modern React)", style='Heading 2')
table(["Hook", "Purpose", "Where in your project"],
      [["useState", "Local component state", "Forms, toggles, every page"],
       ["useEffect", "Side effects & lifecycle", "Fetching, subscriptions"],
       ["useContext", "Read shared context", "AuthContext for the logged-in user"],
       ["useCallback / useMemo", "Memoize functions/values", "useMoneyVisibility, MoneyAmount"],
       ["useRef", "Mutable ref, DOM access", "Speech hooks, media recorder"],
       ["Custom hooks", "Reusable stateful logic", "useSpeechRecognition, useSpeechSynthesis"]],
      widths=[1.5, 2.4, 2.9])
para("Rules of hooks: only call them at the top level of a component (never in loops/conditions) and "
     "only from React functions.", italic=True, color=GREY)

para("2.3 Routing — React Router 7", style='Heading 2')
bullet("BrowserRouter + Routes + Route map URLs to components (see App.jsx).")
bullet("Nested & protected routes: ProtectedRoute checks auth, RoleRoute enforces a role before rendering.")
bullet("Dynamic params: /interview/:token reads the token from the URL for the public candidate portal.")
bullet("Programmatic navigation with useNavigate; redirects with <Navigate />.")

para("2.4 Data fetching — TanStack React Query + Axios", style='Heading 2')
bullet("Axios instance (api/axios.js) sets the base URL and an interceptor that attaches the JWT to every request.")
bullet("React Query manages server state: caching, background refetch, loading/error flags, staleTime.")
bullet("useQuery for reads, useMutation for writes (then invalidate queries to refresh the cache).")
bullet("Why React Query over raw useEffect+fetch: dedupes requests, caches, retries, and removes boilerplate.")

para("2.5 Other frontend skills", style='Heading 2')
bullet("Context API for global state (auth) without prop-drilling.")
bullet("Controlled forms: inputs bound to state, onChange handlers.")
bullet("Conditional rendering, lists, and event handling.")
bullet("CSS: Flexbox, Grid, responsive/mobile-first design, your custom CSS design system + Tailwind config.")
bullet("Browser APIs you actually used: Web Speech API (STT/TTS), MediaRecorder (audio), Page Visibility API + window.blur (proctoring).")
bullet("Vite: fast dev server + bundler; env vars via import.meta.env.VITE_*.")

qa("What is the virtual DOM and why does it help?",
   "It's an in-memory representation of the UI. On state change React builds a new virtual tree, "
   "diffs it against the old one (reconciliation), and applies the minimal set of real DOM updates — "
   "DOM operations are expensive, so this is faster than re-rendering everything.")
qa("Difference between state and props?",
   "Props are passed in from a parent and are read-only; state is owned and mutated by the component "
   "itself. Changing either triggers a re-render.")
qa("Why use React Query instead of fetching in useEffect?",
   "It gives caching, automatic background refetching, request deduplication, retry, and built-in "
   "loading/error states — eliminating a lot of manual, bug-prone state management for server data.")

doc.add_page_break()

# ============================================================ 3. BACKEND
para("3. Backend — Node.js + Express", style='Heading 1')

para("3.1 Node.js", style='Heading 2')
bullet("JavaScript runtime built on V8 that runs outside the browser, with access to the filesystem, network, etc.")
bullet("npm manages dependencies (package.json); scripts run the server (start / dev with nodemon).")
bullet("Non-blocking, event-driven I/O makes it well suited to API servers.")

para("3.2 Express & the request lifecycle", style='Heading 2')
bullet("Express is a minimal web framework: it routes requests and runs middleware.")
bullet("Middleware: functions with (req, res, next) that run in order — CORS, JSON body parsing, auth, error handling. See server.js.")
bullet("Routers split endpoints by domain (routes/*.js), each delegating to a controller.")
bullet("Order matters: candidateRoutes is registered before screeningRoutes so the public token route isn't shadowed.")
bullet("Central error handler (the 4-arg middleware) catches thrown/validation/duplicate-key errors and returns clean JSON.")

para("3.3 Architecture — the MVC-style layering you used", style='Heading 2')
table(["Layer", "Responsibility", "Folder"],
      [["Routes", "Define URLs + attach middleware (auth, RBAC)", "routes/"],
       ["Controllers", "Business logic per domain", "controllers/"],
       ["Models", "Data schema + DB access", "models/"],
       ["Services", "External integrations (Gemini, NN)", "services/"],
       ["Middleware", "Cross-cutting (JWT verify, roles)", "middleware/"],
       ["Utils", "Pure helpers (chunking, tokens, PDF)", "utils/"],
       ["Config", "DB connection, file uploads", "config/"]],
      widths=[1.3, 3.2, 1.3])
para("Why layer it this way: separation of concerns. Each file has one job, so the code is testable, "
     "reusable, and easy to navigate — a point worth making in the interview.", italic=True, color=GREY)

para("3.4 Real-time & file uploads", style='Heading 2')
bullet("Socket.io: persistent WebSocket connection for real-time events (rooms, emit/on). Set up in server.js.")
bullet("Multer: middleware that parses multipart/form-data to accept file uploads (resumes, policy PDFs, audio).")

qa("What is middleware in Express?",
   "A function that receives (req, res, next) and runs during the request–response cycle. It can "
   "inspect/modify req/res, end the response, or call next() to pass control on. I use it for CORS, "
   "body parsing, JWT auth, role checks, and a final error handler.")
qa("How would you structure a Node/Express app and why?",
   "Layered: routes define endpoints and attach middleware, controllers hold business logic, models "
   "define schemas/DB access, services wrap external systems, and utils hold pure helpers. This keeps "
   "concerns separated so each piece is independently testable and reusable.")

doc.add_page_break()

# ============================================================ 4. DATABASE
para("4. Database — MongoDB + Mongoose", style='Heading 1')

para("4.1 MongoDB basics", style='Heading 2')
bullet("NoSQL document database: data stored as flexible JSON-like documents (BSON) in collections.")
bullet("No rigid table schema — good for evolving data, and a natural fit since JS objects map to documents.")
bullet("Documents can embed sub-documents or reference other documents by _id.")

para("4.2 Mongoose ODM", style='Heading 2')
bullet("ODM = Object Data Modeling: defines schemas, validation, and gives a clean query API over MongoDB.")
bullet("Schema → Model → documents. Your models/ folder has 13 schemas (User, Employee, Attendance, etc.).")
bullet("Relationships via ObjectId references + .populate() to join related documents at query time.")
bullet("Common ops: create, find/findById, findByIdAndUpdate, deleteOne, aggregate for stats.")
bullet("Hooks (pre-save) and instance methods — e.g. hashing passwords before save, comparing passwords.")

para("4.3 Data modeling decisions", style='Heading 2')
bullet("Embed vs reference: embed data read together (e.g. interview transcript inside ScreeningSession); reference shared/large data (Employee → User).")
bullet("Why a separate ResumeChunk / RagChunk collection: store many small embedded vectors per document for similarity search.")
bullet("Indexes speed up lookups (e.g. unique index on user email) — be ready to mention them.")

qa("SQL vs NoSQL — why MongoDB here?",
   "SQL databases use fixed tables/rows with strong relational guarantees and joins; NoSQL document "
   "stores like MongoDB use flexible JSON-like documents. I chose MongoDB because the data is "
   "document-shaped, the schema evolved during a hackathon, and JS objects map directly to documents "
   "with Mongoose, speeding development.")
qa("What is populate()?",
   "Mongoose's way of replacing an ObjectId reference with the actual referenced document at query "
   "time — effectively a join, e.g. loading the User behind an Employee.")
qa("How do you secure passwords?",
   "Never store plaintext. I hash with bcrypt (a slow, salted hashing algorithm) in a pre-save hook, "
   "and compare on login with bcrypt.compare. Salting + slowness defends against rainbow tables and "
   "brute force.")

doc.add_page_break()

# ============================================================ 5. AUTH & SECURITY
para("5. Authentication, Authorization & Security", style='Heading 1')

para("5.1 The auth flow in your app", style='Heading 2')
bullet("Login: verify email + bcrypt-compared password → sign a JWT containing the user id.")
bullet("Token storage: frontend keeps the JWT; Axios interceptor adds 'Authorization: Bearer <token>' to every request.")
bullet("protect middleware: verifies the JWT signature, loads the user, attaches req.user.")
bullet("authorize(...roles) middleware: checks req.user.role against allowed roles (RBAC).")
bullet("Public candidate portal: no JWT — access is gated by a per-session token in the URL with a 48h expiry.")

para("5.2 Concepts to be able to explain", style='Heading 2')
bullet("Authentication (who you are) vs authorization (what you may do).", bold_lead="")
bullet("JWT structure: header.payload.signature — signed (not encrypted) with a secret; tamper-evident.")
bullet("Why JWT: stateless auth, no server session store needed; the token carries the claims.")
bullet("bcrypt: salted, slow hashing; salt prevents identical passwords from hashing alike.")
bullet("RBAC: role-based access control — permissions attached to roles, enforced per route.")
bullet("Common risks & mitigations: never commit secrets (.env), validate input, HTTPS in prod, set token expiry, restrict CORS.")

qa("How does JWT-based auth work end to end?",
   "On login the server verifies credentials and signs a JWT (user id in the payload) with a secret. "
   "The client stores it and sends it in the Authorization header. Middleware verifies the signature "
   "on each request; if valid it trusts the claims and attaches the user. Roles are then checked for "
   "authorization. It's stateless — no session storage server-side.")
qa("JWT vs sessions?",
   "Sessions store state server-side and hand the client a session id; JWTs are self-contained and "
   "stateless, so they scale horizontally without a shared session store — at the cost of being hard "
   "to revoke before expiry.")
qa("How do you protect a public endpoint like the candidate interview?",
   "Instead of a login, I issue a random, time-limited token tied to that one screening session and "
   "embed it in the URL. The server validates the token and its 48-hour expiry on every candidate "
   "request, so only the intended candidate can access just their own session.")

doc.add_page_break()

# ============================================================ 6. AI / ML
para("6. AI & Machine Learning Integration", style='Heading 1')
para("This is the differentiator of your project — and where most interview questions will focus. "
     "Be able to explain each feature, the technique behind it, and the trade-offs.")

para("6.1 Large Language Models (LLMs) & the Gemini API", style='Heading 2')
bullet("An LLM predicts text from a prompt; you call it over an API (Google Gemini) with a prompt and get generated text back.")
bullet("Prompt engineering: clear instructions, context, role, and output format drive quality.")
bullet("Tokens: text is split into tokens; cost/limits are token-based. Temperature controls randomness.")
bullet("Your fallback chain (geminiService.js): try gemini-2.5-flash → -flash-lite → 2.0-flash → 1.5-flash, "
       "catching 404/429/503 with back-off. This is resilience engineering against rate limits/outages — a strong talking point.")

para("6.2 Embeddings & RAG (Retrieval-Augmented Generation)", style='Heading 2')
para("RAG is how your HR policy chatbot answers from company documents instead of hallucinating.")
bullet("Embedding: a model turns text into a vector (list of numbers) capturing meaning. Similar meaning → nearby vectors.")
bullet("Pipeline (ingest): PDF → extract text → chunk into overlapping pieces (chunkText.js) → embed each chunk → store vectors in MongoDB (RagChunk).")
bullet("Pipeline (query): embed the user's question → cosine-similarity search to find the top-k closest chunks → inject them into the Gemini prompt as context → Gemini answers grounded in those chunks.")
bullet("Why RAG: keeps answers accurate and current without retraining the model; reduces hallucination; cites real policy text.")
bullet("Cosine similarity: measures the angle between two vectors (1 = identical direction). The core of the retrieval step.")

para("6.3 The neural network — flight-risk prediction (Brain.js)", style='Heading 2')
bullet("Goal: predict an employee's churn risk (low/medium/high/critical) from attendance, performance, and leave features.")
bullet("Brain.js: a JS neural network library. You train on labeled data and save the weights (ml_models/flight_risk_nn.json).")
bullet("Features: ~8 normalized numeric inputs → network → a single risk score 0–1, bucketed into levels.")
bullet("Training data: generated synthetically (generateSyntheticData.js) since real churn labels weren't available — be honest about this trade-off.")
bullet("Key ML vocabulary: features, labels, training vs inference, normalization, overfitting, weights, activation.")

para("6.4 AI interviews + proctoring", style='Heading 2')
bullet("Adaptive interview: Gemini conducts up to 8 turns, generating follow-up questions from the candidate's answers (interviewHelpers.js manages turns + transcript).")
bullet("Voice: candidate speech recorded via MediaRecorder → sent to server → transcribed by Gemini (speech-to-text) → returned as text.")
bullet("Proctoring: Page Visibility API + window.blur detect tab-switching/focus loss; after 3 violations the session is flagged. Events are logged silently.")
bullet("Analysis: at the end, Gemini scores communication, technical depth, and confidence and produces a hire/no-hire verdict.")

para("6.5 Responsible-AI angle (great to raise unprompted)", style='Heading 2')
bullet("Blind screening: PII is hidden until shortlisting to reduce unconscious bias.")
bullet("Limitations: LLMs can hallucinate (mitigated by RAG); synthetic training data limits the NN's real-world accuracy; AI assists humans, it doesn't replace the final decision.")

qa("What is RAG and why did you use it?",
   "Retrieval-Augmented Generation: instead of relying on the LLM's parametric memory, I retrieve "
   "relevant document chunks and feed them to the model as context. I embed the policy PDFs into "
   "vectors, store them in MongoDB, embed the question at query time, find the most similar chunks by "
   "cosine similarity, and inject them into the prompt. This grounds answers in real company policy "
   "and sharply reduces hallucination, with no model retraining.")
qa("What is an embedding?",
   "A numeric vector representation of text where semantic similarity corresponds to geometric "
   "closeness. It lets me compare meaning mathematically — the basis of semantic search in RAG.")
qa("Why a fallback chain of models?",
   "Free-tier LLM APIs hit rate limits (429) and occasional outages (503). My service tries a "
   "preferred model first and degrades gracefully to alternatives with back-off, so the feature stays "
   "available instead of failing on a single error.")
qa("How does your flight-risk model work and what are its limits?",
   "A small Brain.js neural net takes ~8 normalized features from attendance, performance and leave "
   "data and outputs a 0–1 risk score bucketed into four levels. Because I lacked real attrition "
   "labels, I trained on synthetic data — so it demonstrates the pipeline and is directionally useful, "
   "but I'd retrain on real historical outcomes before trusting it operationally.")
qa("How do you prevent the LLM from leaking bias into hiring?",
   "Blind screening hides candidate PII until shortlisting, the model is evaluated against the job "
   "description rather than identity, and AI output is a recommendation a human reviews — it never "
   "auto-rejects.")

doc.add_page_break()

# ============================================================ 7. TOOLING & DEPLOYMENT
para("7. Tooling, DevOps & Deployment", style='Heading 1')
bullet("Git & GitHub: branches, commits, pull requests, resolving conflicts — essential and often asked.")
bullet("Environment variables: secrets and config in .env (never committed); separate dev vs production values.")
bullet("npm: dependencies vs devDependencies, lockfiles, semantic versioning.")
bullet("Vite build for the frontend; nodemon for backend hot-reload in dev.")
bullet("Deployment: render.yaml provisions a backend web service + frontend static site on Render.com; "
       "MongoDB Atlas for the cloud database; set env vars in the dashboard; vercel.json handles SPA routing if hosted on Vercel.")
bullet("Production hardening you implemented: server refuses to boot if CLIENT_URL is misconfigured in production.")
bullet("Debugging skills: reading stack traces, console/network tab, Postman/REST client (rsq.http), checking logs.")

qa("Walk me through how this app is deployed.",
   "The repo is a monorepo. Render reads render.yaml and provisions two services: the Express API as a "
   "web service and the built React app as a static site. The database is MongoDB Atlas. Secrets "
   "(Mongo URI, JWT secret, Gemini key, client URL) are set in the Render dashboard, not in the repo. "
   "After deploy I seed the DB via the provided scripts.")

doc.add_page_break()

# ============================================================ 8. SYSTEM DESIGN
para("8. System Design — Tell the Whole Story", style='Heading 1')
para("Interviewers love a clean end-to-end narrative. Practice saying this out loud:")
para("\"A user logs in; the React SPA stores a JWT and attaches it to every Axios request. Requests hit "
     "the Express API, where middleware verifies the JWT and the user's role before routing to a "
     "controller. Controllers run business logic against MongoDB through Mongoose models. For AI "
     "features, controllers call a Gemini service (with a model fallback chain) or a Brain.js neural "
     "net. The policy chatbot uses RAG: documents are chunked, embedded, and stored as vectors; at "
     "query time the closest chunks are retrieved by cosine similarity and injected into the prompt. "
     "Candidates take AI interviews through a public, token-gated link with browser-based proctoring.\"",
     italic=True)

para("8.1 Trade-offs you should be ready to discuss", style='Heading 2')
table(["Decision", "Why", "Trade-off / alternative"],
      [["MongoDB (NoSQL)", "Flexible, document-shaped data, fast iteration", "Weaker for complex relational queries vs SQL"],
       ["JWT stateless auth", "Scales without session store", "Hard to revoke before expiry"],
       ["RAG over fine-tuning", "Accurate, current, no retraining", "Needs a vector store + retrieval step"],
       ["Synthetic NN data", "No real labels available", "Limited real-world accuracy"],
       ["Gemini free tier", "Zero cost for a hackathon", "Rate limits → needed a fallback chain"],
       ["Custom CSS (no UI lib)", "Full control, no bloat", "More work than a component library"]],
      widths=[1.6, 2.7, 2.4])

para("8.2 \"How would you improve / scale it?\" (have answers ready)", style='Heading 2')
bullet("Move vector search to a dedicated vector DB (e.g. Atlas Vector Search / Pinecone) for speed at scale.")
bullet("Add automated tests (Jest, React Testing Library) and CI.")
bullet("Add caching (Redis) and rate limiting; paginate large lists.")
bullet("Refresh tokens + httpOnly cookies for stronger auth than localStorage.")
bullet("Train the flight-risk model on real historical attrition data; add monitoring for drift.")
bullet("Observability: structured logging, error tracking (Sentry), health metrics.")

doc.add_page_break()

# ============================================================ 9. RAPID-FIRE
para("9. Rapid-Fire Question Bank", style='Heading 1')
para("Be able to answer each in 1–2 sentences.", italic=True, color=GREY)

para("JavaScript / React", style='Heading 3')
for q in [
    "What is the event loop?",
    "Explain closures with an example.",
    "var vs let vs const?",
    "What does useEffect's dependency array do?",
    "Why do React lists need keys?",
    "Controlled vs uncontrolled components?",
    "What is prop drilling and how does Context solve it?",
    "What is a custom hook?",
]:
    bullet(q)

para("Backend / Database", style='Heading 3')
for q in [
    "What is middleware? Give the execution order in your app.",
    "How does Express handle errors centrally?",
    "Embedding vs referencing in MongoDB?",
    "What is an index and when do you add one?",
    "What does Mongoose's populate do?",
    "How do you hash and verify passwords?",
    "Why register the candidate route before the screening route?",
]:
    bullet(q)

para("AI / ML", style='Heading 3')
for q in [
    "What is a token / what is temperature?",
    "What is prompt engineering?",
    "Explain RAG in 3 sentences.",
    "What is cosine similarity?",
    "Difference between training and inference?",
    "What is overfitting and how do you reduce it?",
    "Why might an LLM hallucinate and how do you mitigate it?",
    "What features feed your flight-risk model?",
]:
    bullet(q)

para("Project-specific", style='Heading 3')
for q in [
    "Walk me through the architecture end to end.",
    "What was the hardest part to build and how did you solve it?",
    "How do the four roles differ in access?",
    "How does proctoring detect cheating?",
    "What would you do differently with more time?",
    "How does the blind screening reduce bias?",
]:
    bullet(q)

doc.add_page_break()

# ============================================================ 10. STUDY PLAN
para("10. A Focused 2-Week Study Plan", style='Heading 1')
table(["Days", "Focus", "Outcome"],
      [["1–2", "JS fundamentals + async/event loop", "Explain closures, promises, the event loop"],
       ["3–4", "React: components, hooks, routing, React Query", "Rebuild one of your pages from memory"],
       ["5–6", "Node + Express: middleware, MVC, error handling", "Trace a request from route → controller → DB"],
       ["7", "MongoDB + Mongoose: schemas, populate, modeling", "Justify your schema choices"],
       ["8", "Auth & security: JWT, bcrypt, RBAC", "Explain the full login → protected request flow"],
       ["9–10", "AI: LLMs, prompts, embeddings, RAG", "Explain the policy chatbot pipeline cold"],
       ["11", "ML: neural net basics, your flight-risk model", "Explain features, training, limits"],
       ["12", "Deployment, Git, env config", "Describe how it ships to production"],
       ["13", "System design narrative + trade-offs", "Deliver the end-to-end story smoothly"],
       ["14", "Mock interview with the question bank", "Answer rapid-fire without notes"]],
      widths=[0.8, 3.0, 2.9])

para("Final advice", style='Heading 2')
bullet("Build, don't just read. Re-implement a slice (e.g. login + one CRUD module) from scratch.")
bullet("Know your own code: open each folder and be ready to explain why it exists.")
bullet("Be honest about trade-offs and limitations — interviewers respect it more than overclaiming.")
bullet("Tell stories: 'I hit X problem, tried Y, and solved it with Z' is far stronger than reciting definitions.")

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("Good luck — you built it, so you can explain it.")
r.italic = True; r.font.color.rgb = ACCENT; r.font.size = Pt(12)

doc.save(r"d:\Documents\FWC\FWC_HRMS_Interview_Prep_Guide.docx")
print("Saved FWC_HRMS_Interview_Prep_Guide.docx")
